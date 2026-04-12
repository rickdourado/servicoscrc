from fastapi import FastAPI, File, UploadFile
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from typing import List
import backend.scripts.core_logic as core_logic
from pydantic import BaseModel
import os
from google import genai
from dotenv import load_dotenv
import backend.scripts.anonymizer as anonymizer
from pypdf import PdfReader
import io

load_dotenv()

app = FastAPI(title="App Lúdico API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

os.makedirs(anonymizer.TEMP_DIR, exist_ok=True)
app.mount("/temp", StaticFiles(directory=anonymizer.TEMP_DIR), name="temp")

class OrderData(BaseModel):
    items: List[core_logic.Level1Item]

class AnalyzeTextData(BaseModel):
    text: str
    filename: str | None = None

@app.get("/api/data")
def get_data():
    items = core_logic.extract_data()
    return {"items": items}

@app.post("/api/save")
def save_data(data: OrderData):
    result = core_logic.save_new_order(data.items)
    return result

@app.post("/api/analyze-contract")
async def analyze_contract(file: UploadFile = File(...)):
    try:
        with open("backend/prompts/analise_contrato.txt", "r", encoding="utf-8") as f:
            prompt = f.read()
    except Exception as e:
        return {"result": f"Erro abrindo o arquivo de prompt: {str(e)}"}

    file_bytes = await file.read()
    
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return {"result": f"MOCK RESULT (Modo Offline - Key Inexistente):\n\n🚨 O arquivo '{file.filename}' (Tamanho: ~{len(file_bytes)//1024} KB) foi recebido com sucesso!\n\nNo entanto, o backend não detectou a variável GEMINI_API_KEY.\nO Prompt que seria passado para a inteligência é:\n\n\"{prompt}\"\n\n➡️ Defina a variável 'GEMINI_API_KEY' no sistema antes do 'uv run' nas próximas análises."}
    
    try:
        client = genai.Client(api_key=api_key)
        
        raw_text = ""
        is_pdf = "pdf" in (file.content_type or "").lower() or file.filename.lower().endswith(".pdf")
        
        if is_pdf:
             reader = PdfReader(io.BytesIO(file_bytes))
             for page in reader.pages:
                 text_page = page.extract_text()
                 if text_page:
                     raw_text += text_page + "\n"
        else:
             try:
                 import docx
                 doc = docx.Document(io.BytesIO(file_bytes))
                 for para in doc.paragraphs:
                     raw_text += para.text + "\n"
             except Exception:
                 raw_text = file_bytes.decode("utf-8", errors="ignore")

        # Aplica a censura de PII localmente
        masked_text = anonymizer.process_and_save(raw_text, file.filename)
        
        # Envia o prompt modificado e apenas a RAW string já mascarada (sem bytes de arquivo)
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[prompt + "\n\n(Aviso do sistema interno: o conteúdo a seguir foi lido localmente e nós ocultamos propositalmente os nomes reais da empresa e CNPJ substituindo por tags [OCULTO]. Aceite e considere como parte normal do documento.)\n\n--- INÍCIO DO TEXTO ---",
            masked_text]
        )
             
        return {"result": response.text}
    except Exception as e:
        return {"result": f"Erro interno ao invocar Inteligência Artificial: {str(e)}"}

@app.post("/api/anonymize")
async def anonymize_contract(file: UploadFile = File(...)):
    file_bytes = await file.read()
    raw_text = ""
    is_pdf = "pdf" in (file.content_type or "").lower() or file.filename.lower().endswith(".pdf")
    
    if is_pdf:
        filename, masked_text = anonymizer.redact_pdf_visually(file_bytes)
        return {"masked_text": masked_text, "preview_url": f"/temp/{filename}"}
    else:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                raw_text += para.text + "\n"
        except Exception:
            raw_text = file_bytes.decode("utf-8", errors="ignore")

        # Aplica a censura de PII localmente via raw text
        masked_text = anonymizer.process_and_save(raw_text, file.filename)
        return {"masked_text": masked_text, "preview_url": None}

@app.post("/api/analyze-text")
async def analyze_text(data: AnalyzeTextData):
    try:
        with open("backend/prompts/analise_contrato.txt", "r", encoding="utf-8") as f:
            prompt = f.read()
    except Exception as e:
        return {"result": f"Erro abrindo o arquivo de prompt: {str(e)}"}

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return {"result": f"MOCK RESULT (Modo Offline - Key Inexistente):\n\n🚨 O texto foi recebido com sucesso!\n\nNo entanto, o backend não detectou a variável GEMINI_API_KEY.\nO Prompt que seria passado para a inteligência é:\n\n\"{prompt}\"\n\n➡️ Defina a variável 'GEMINI_API_KEY' no sistema antes do 'uv run' nas próximas análises."}
    
    try:
        client = genai.Client(api_key=api_key)
        
        contents_list = [prompt + "\n\n(Aviso do sistema interno: o conteúdo a seguir foi lido localmente e nós ocultamos propositalmente os nomes reais da empresa e CNPJ substituindo por tarjas pretas de redação [OCULTO]. Considere essas tarjas apenas como proteção e analise o layout da página como parte normal do documento.)\n\n--- INÍCIO DO DOCUMENTO ---"]
        
        if data.filename and data.filename.endswith(".pdf"):
            pdf_path = os.path.join(anonymizer.TEMP_DIR, data.filename)
            if os.path.exists(pdf_path):
                uploaded_file = client.files.upload(file=pdf_path)
                contents_list.append(uploaded_file)
            else:
                contents_list.append(data.text)
        else:
            contents_list.append(data.text)
            
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=contents_list
        )
             
        return {"result": response.text}
    except Exception as e:
        return {"result": f"Erro interno ao invocar Inteligência Artificial: {str(e)}"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
